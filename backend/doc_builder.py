"""
Document Builder Module
Builds the output Word (.docx) document matching Document A's exact layout:
Center-aligned Main Titles & Subtitles, Left-aligned Section Headings,
and 100% strictly the content and blank spaces that exist in Document B.
"""
import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from style_extractor import StyleTokens
from content_parser import ParsedDocument, ContentBlock

def hex_to_rgb(hex_str: str) -> RGBColor:
    hex_clean = hex_str.lstrip('#')
    if len(hex_clean) == 6:
        r = int(hex_clean[0:2], 16)
        g = int(hex_clean[2:4], 16)
        b = int(hex_clean[4:6], 16)
        return RGBColor(r, g, b)
    return RGBColor(0x1F, 0x37, 0x64)


def build_styled_document(
    template_docx_path: str,
    parsed_doc: ParsedDocument,
    styles: StyleTokens,
    output_docx_path: str
) -> str:
    """Builds a styled Word document rendering exclusively Document B data with Document A style."""
    
    # 1. Start from base template if available to preserve styles
    if os.path.exists(template_docx_path) and template_docx_path.lower().endswith('.docx'):
        shutil.copy2(template_docx_path, output_docx_path)
        doc = Document(output_docx_path)
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('p', 'tbl'):
                body.remove(child)
    else:
        doc = Document()

    # Set page margins matching Document A (top: 1080 dxa, bottom: 500 dxa, left: 992 dxa, right: 992 dxa)
    for section in doc.sections:
        section.top_margin = Inches(1080 / 1440.0)
        section.bottom_margin = Inches(500 / 1440.0)
        section.left_margin = Inches(992 / 1440.0)
        section.right_margin = Inches(992 / 1440.0)

    primary_rgb = hex_to_rgb(styles.primary_color)
    
    # 2. Iterate through parsed blocks and render each
    for block_idx, block in enumerate(parsed_doc.blocks):
        b_type = block.block_type
        text = block.text
        
        if b_type == 'title':
            # Main Course Title (CENTER ALIGNED, Heading 1 in primary accent color)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text.upper())
            run.bold = True
            run.font.name = styles.font_family
            run.font.size = Pt(15.5)
            run.font.color.rgb = primary_rgb

        elif b_type == 'subtitle':
            # Subtitle (CENTER ALIGNED, Italic)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.italic = True
            run.font.name = styles.font_family
            run.font.size = Pt(11)

        elif b_type == 'heading':
            # Section Heading (LEFT ALIGNED in primary accent bold)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = styles.font_family
            run.font.size = Pt(styles.heading3_size)
            run.font.color.rgb = primary_rgb

        elif b_type == 'subheading':
            # Sub-heading (LEFT ALIGNED in primary accent bold)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = styles.font_family
            run.font.size = Pt(styles.subheading_size)
            run.font.color.rgb = primary_rgb

        elif b_type == 'body':
            # Standard Body Paragraph (LEFT ALIGNED, NO BULLET)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(styles.body_space_before)
            p.paragraph_format.space_after = Pt(styles.body_space_after)
            p.paragraph_format.left_indent = Pt(6)
            p.paragraph_format.line_spacing = styles.line_spacing
            run = p.add_run(text)
            run.font.name = styles.font_family
            run.font.size = Pt(styles.body_size)

        elif b_type == 'bullet':
            # Bullet point with hanging indent (LEFT ALIGNED)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(styles.bullet_space_before)
            p.paragraph_format.space_after = Pt(styles.bullet_space_after)
            p.paragraph_format.left_indent = Pt(22)
            p.paragraph_format.first_line_indent = Pt(-14)
            p.paragraph_format.line_spacing = styles.line_spacing
            
            run_bullet = p.add_run("• ")
            run_bullet.bold = True
            run_bullet.font.name = styles.font_family
            run_bullet.font.size = Pt(styles.body_size)
            run_bullet.font.color.rgb = primary_rgb
            
            run_text = p.add_run(text)
            run_text.font.name = styles.font_family
            run_text.font.size = Pt(styles.body_size)

        elif b_type == 'numbered':
            # Numbered list item with hanging indent (LEFT ALIGNED)
            num_str = block.extra.get('num', '1.')
            if not num_str.endswith(('.', ')', ':')):
                num_str = f"{num_str}."
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(styles.bullet_space_before)
            p.paragraph_format.space_after = Pt(styles.bullet_space_after)
            p.paragraph_format.left_indent = Pt(26)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.line_spacing = styles.line_spacing
            
            run_num = p.add_run(f"{num_str} ")
            run_num.bold = True
            run_num.font.name = styles.font_family
            run_num.font.size = Pt(styles.body_size)
            run_num.font.color.rgb = primary_rgb
            
            run_text = p.add_run(text)
            run_text.font.name = styles.font_family
            run_text.font.size = Pt(styles.body_size)

        elif b_type == 'alpha':
            # Alphabetical list item with hanging indent (LEFT ALIGNED)
            alpha_str = block.extra.get('alpha', 'a.')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(styles.bullet_space_before)
            p.paragraph_format.space_after = Pt(styles.bullet_space_after)
            p.paragraph_format.left_indent = Pt(26)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.line_spacing = styles.line_spacing
            
            run_a = p.add_run(f"{alpha_str} ")
            run_a.bold = True
            run_a.font.name = styles.font_family
            run_a.font.size = Pt(styles.body_size)
            run_a.font.color.rgb = primary_rgb
            
            run_text = p.add_run(text)
            run_text.font.name = styles.font_family
            run_text.font.size = Pt(styles.body_size)

        elif b_type == 'outcome':
            # Course Outcome (LEFT ALIGNED with hanging indent)
            co_id = block.extra.get('co_id', 'CO')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.line_spacing = styles.line_spacing
            
            run_id = p.add_run(f"{co_id} ")
            run_id.bold = True
            run_id.font.name = styles.font_family
            run_id.font.size = Pt(styles.body_size)
            run_id.font.color.rgb = primary_rgb
            
            run_text = p.add_run(text)
            run_text.font.name = styles.font_family
            run_text.font.size = Pt(styles.body_size)

        elif b_type == 'empty_space':
            # Intentional blank space / empty line preserved from Document B
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(10)

        elif b_type == 'table':
            headers = block.extra.get('headers', [])
            rows = block.extra.get('rows', [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Repeat header row across multiple pages
                header_tr = table.rows[0]._tr.get_or_add_trPr()
                header_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                
                # Header row formatting
                for i, title in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = title
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>')
                    tcPr.append(tcMar)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or len(title.split()) <= 2) else WD_ALIGN_PARAGRAPH.LEFT
                        for r in p.runs:
                            r.bold = True
                            r.font.name = styles.font_family
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            r.font.size = Pt(styles.table_head_size)
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{styles.table_header_fill}"/>')
                    tcPr.append(shd)

                # Data rows formatting
                for r_idx, row_vals in enumerate(rows):
                    trPr = table.rows[r_idx + 1]._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                    for c_idx, val in enumerate(row_vals):
                        if c_idx < len(table.rows[r_idx + 1].cells):
                            cell = table.rows[r_idx + 1].cells[c_idx]
                            cell.text = str(val)
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="70" w:type="dxa"/><w:bottom w:w="70" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>')
                            tcPr.append(tcMar)
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (c_idx == 0 and len(headers) > 4) else WD_ALIGN_PARAGRAPH.LEFT
                                for r in p.runs:
                                    r.font.name = styles.font_family
                                    r.font.size = Pt(styles.table_text_size)

                tblPr = table._tbl.tblPr
                borders = parse_xml(
                    f'<w:tblBorders {nsdecls("w")}>'
                    f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{styles.table_border_color}"/>'
                    f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="{styles.table_border_color}"/>'
                    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{styles.table_border_color}"/>'
                    f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="{styles.table_border_color}"/>'
                    f'  <w:insideH w:val="single" w:sz="8" w:space="0" w:color="{styles.table_border_color}"/>'
                    f'  <w:insideV w:val="single" w:sz="8" w:space="0" w:color="{styles.table_border_color}"/>'
                    '</w:tblBorders>'
                )
                tblPr.append(borders)

        elif b_type == 'page_break':
            doc.add_page_break()

    # Save document
    doc.save(output_docx_path)
    return output_docx_path
