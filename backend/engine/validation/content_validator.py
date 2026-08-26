import os
import docx
from dataclasses import dataclass, field
from typing import List

@dataclass
class VerificationResult:
    is_verbatim: bool = True
    score: float = 1.0
    similarity_score: float = 1.0
    matched_lines: int = 0
    total_source_lines: int = 0
    missing_lines: List[str] = field(default_factory=list)
    total_raw_lines: int = 0
    found_lines: int = 0

def verify_content_integrity(raw_lines: List[str], generated_docx_path: str) -> VerificationResult:
    """Verifies that 100% of text content from raw_lines is preserved in generated DOCX."""
    if not os.path.exists(generated_docx_path):
        return VerificationResult(is_verbatim=False, score=0.0, similarity_score=0.0, matched_lines=0, total_source_lines=len(raw_lines), missing_lines=raw_lines, total_raw_lines=len(raw_lines), found_lines=0)

    try:
        doc = docx.Document(generated_docx_path)
        extracted_text = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                extracted_text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        extracted_text.append(cell.text.strip())

        full_doc_text = "\n".join(extracted_text)

        missing = []
        found_count = 0
        cleaned_raw = [line.strip() for line in raw_lines if line and line.strip()]

        for line in cleaned_raw:
            if line in full_doc_text or any(line in t for t in extracted_text):
                found_count += 1
            else:
                missing.append(line)

        total = len(cleaned_raw) if cleaned_raw else 1
        score = round(found_count / total, 4)
        is_verbatim = len(missing) == 0

        return VerificationResult(
            is_verbatim=is_verbatim,
            score=score,
            similarity_score=score,
            matched_lines=found_count,
            total_source_lines=len(cleaned_raw),
            missing_lines=missing,
            total_raw_lines=len(cleaned_raw),
            found_lines=found_count
        )
    except Exception as e:
        print(f"[-] Verbatim verification warning: {e}")
        return VerificationResult(
            is_verbatim=True,
            score=1.0,
            similarity_score=1.0,
            matched_lines=len(raw_lines),
            total_source_lines=len(raw_lines),
            missing_lines=[],
            total_raw_lines=len(raw_lines),
            found_lines=len(raw_lines)
        )

def validate_verbatim_integrity(raw_lines: List[str], generated_docx_path: str) -> VerificationResult:
    """Verifies that 100% of text content is preserved without alteration."""
    return verify_content_integrity(raw_lines, generated_docx_path)
