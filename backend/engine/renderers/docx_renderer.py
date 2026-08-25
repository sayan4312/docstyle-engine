import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from ..models.document_ast import CanonicalAST
from ..models.template_model import TemplateModel, ElementStyle
from ..models.semantic_block import SemanticType

def hex_to_rgb(hex_str: str) -> RGBColor:
    hex_clean = hex_str.lstrip('#') if hex_str else "000000"
    if len(hex_clean) != 6:
        hex_clean = "000000"
    try:
        r = int(hex_clean[0:2], 16)
        g = int(hex_clean[2:4], 16)
        b = int(hex_clean[4:6], 16)
        return RGBColor(r, g, b)
    except Exception:
        return RGBColor(0, 0, 0)


def render_ast_to_docx(ast: CanonicalAST, template: TemplateModel, output_path: str) -> str:
    """Renders CanonicalAST and TemplateModel into a formatted Word .docx file."""
    doc = docx.Document()
    
    # 1. Apply Page Margins
    section = doc.sections[0]
    section.top_margin = Pt(template.margin_top / 20)
    section.bottom_margin = Pt(template.margin_bottom / 20)
    section.left_margin = Pt(template.margin_left / 20)
    section.right_margin = Pt(template.margin_right / 20)
    section.page_width = Pt(template.page_width / 20)
    section.page_height = Pt(template.page_height / 20)

    # 2. Default Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = template.styles.get("PARAGRAPH", ElementStyle()).font_family
    style_normal.font.size = Pt(template.styles.get("PARAGRAPH", ElementStyle()).font_size)

    primary_rgb = hex_to_rgb(template.primary_color)
    secondary_rgb = hex_to_rgb(template.secondary_color)
    body_rgb = hex_to_rgb(template.body_text_color)

    # Track rendered non-empty blocks to apply page break on non-first TITLEs
    rendered_non_empty_count = 0

    # 3. Iterate over AST blocks and render elements
    for block in ast.blocks:
        b_type = block.type
        text = block.text
        override = block.style_override or {}
        font_name = override.get("font_family", template.styles.get("PARAGRAPH", ElementStyle()).font_family)

        if b_type == SemanticType.TITLE.value:
            # Force page break for subsequent major Titles starting mid-document
            if rendered_non_empty_count > 0:
                doc.add_page_break()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(override.get("space_before", 16.0))
            p.paragraph_format.space_after = Pt(override.get("space_after", 4.0))
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(override.get("font_size", 20.0))
            run.font.color.rgb = primary_rgb
            rendered_non_empty_count += 1

        elif b_type == SemanticType.SUBTITLE.value:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2.0)
            p.paragraph_format.space_after = Pt(14.0)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(override.get("font_size", 12.0))
            run.font.color.rgb = secondary_rgb
            rendered_non_empty_count += 1

        elif b_type in (SemanticType.HEADING_1.value, SemanticType.HEADING_2.value, SemanticType.HEADING_3.value):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if b_type == SemanticType.HEADING_1.value:
                p.paragraph_format.space_before = Pt(14.0)
                p.paragraph_format.space_after = Pt(4.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(override.get("font_size", 15.0))
                run.font.color.rgb = primary_rgb
            elif b_type == SemanticType.HEADING_2.value:
                p.paragraph_format.space_before = Pt(10.0)
                p.paragraph_format.space_after = Pt(3.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(override.get("font_size", 13.0))
                run.font.color.rgb = secondary_rgb
            else:
                p.paragraph_format.space_before = Pt(8.0)
                p.paragraph_format.space_after = Pt(2.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(override.get("font_size", 11.5))
                run.font.color.rgb = body_rgb
            rendered_non_empty_count += 1

        elif b_type == SemanticType.BULLET_LIST.value:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3.0)
            p.paragraph_format.space_after = Pt(2.0)
            p.paragraph_format.left_indent = Pt(22)
            p.paragraph_format.first_line_indent = Pt(-14)
            p.paragraph_format.line_spacing = override.get("line_spacing", 1.15)

            rb = p.add_run("• ")
            rb.bold = True
            rb.font.name = font_name
            rb.font.size = Pt(override.get("font_size", 10.5))
            rb.font.color.rgb = primary_rgb

            rt = p.add_run(text)
            rt.font.name = font_name
            rt.font.size = Pt(override.get("font_size", 10.5))
            rt.font.color.rgb = body_rgb
            rendered_non_empty_count += 1

        elif b_type == SemanticType.NUMBERED_LIST.value:
            num_str = block.extra.get("num", block.extra.get("alpha", "1."))
            if not num_str.endswith('.'):
                num_str += "."

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3.0)
            p.paragraph_format.space_after = Pt(2.0)
            p.paragraph_format.left_indent = Pt(26)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.line_spacing = override.get("line_spacing", 1.15)

            rn = p.add_run(f"{num_str} ")
            rn.bold = True
            rn.font.name = font_name
            rn.font.size = Pt(override.get("font_size", 10.5))
            rn.font.color.rgb = primary_rgb

            rt = p.add_run(text)
            rt.font.name = font_name
            rt.font.size = Pt(override.get("font_size", 10.5))
            rt.font.color.rgb = body_rgb
            rendered_non_empty_count += 1

        elif b_type == SemanticType.TABLE.value:
            headers = block.extra.get("headers", [])
            rows = block.extra.get("rows", [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                header_tr = table.rows[0]._tr.get_or_add_trPr()
                header_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

                for i, title in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = title
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>')
                    tcPr.append(tcMar)
                    for p_cell in cell.paragraphs:
                        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or len(title.split()) <= 2) else WD_ALIGN_PARAGRAPH.LEFT
                        for r in p_cell.runs:
                            r.bold = True
                            r.font.name = font_name
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            r.font.size = Pt(9.5)
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{template.table_header_fill}"/>')
                    tcPr.append(shd)

                for r_idx, row_vals in enumerate(rows):
                    trPr = table.rows[r_idx + 1]._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                    for c_idx, val in enumerate(row_vals):
                        if c_idx < len(table.rows[r_idx + 1].cells):
                            cell = table.rows[r_idx + 1].cells[c_idx]
                            cell.text = str(val)
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="70" w:type="dxa"/><w:bottom w:w="70" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>')
                            tcPr.append(tcMar)
                            for p_cell in cell.paragraphs:
                                for r in p_cell.runs:
                                    r.font.name = font_name
                                    r.font.size = Pt(9.0)
                                    r.font.color.rgb = body_rgb

                tblPr = table._tbl.tblPr
                borders = parse_xml(
                    f'<w:tblBorders {nsdecls("w")}>'
                    f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{template.table_border_color}"/>'
                    f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="{template.table_border_color}"/>'
                    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{template.table_border_color}"/>'
                    f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="{template.table_border_color}"/>'
                    f'  <w:insideH w:val="single" w:sz="8" w:space="0" w:color="{template.table_border_color}"/>'
                    f'  <w:insideV w:val="single" w:sz="8" w:space="0" w:color="{template.table_border_color}"/>'
                    '</w:tblBorders>'
                )
                tblPr.append(borders)
                rendered_non_empty_count += 1

        elif b_type == SemanticType.PAGE_BREAK.value:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(10)

        else:
            # Paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(override.get("space_before", 4.0))
            p.paragraph_format.space_after = Pt(override.get("space_after", 3.0))
            p.paragraph_format.line_spacing = override.get("line_spacing", 1.15)

            rt = p.add_run(text)
            rt.font.name = font_name
            rt.font.size = Pt(override.get("font_size", 10.5))
            rt.font.color.rgb = body_rgb
            rendered_non_empty_count += 1

    doc.save(output_path)
    return output_path
